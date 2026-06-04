#!/usr/bin/env python3
"""Extract, map, audit, profile, render, and rewrite document text.

The DOCX paths use only Python's standard library so the helper stays usable in
minimal agent environments. PDF support is opportunistic: it uses whichever
common extraction library is already installed. Rebuilt DOCX output uses
python-docx because generating a valid Word package by hand is not worth the
fragility.
"""

from __future__ import annotations

import argparse
from collections import Counter
import html
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path
from statistics import mean, pstdev
from typing import Any, Iterable
from xml.etree import ElementTree as ET

try:
    from lxml import etree as LET  # type: ignore
except ImportError:  # pragma: no cover - validated at runtime for DOCX writes.
    LET = None


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
W_NS = NS["w"]
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
WORD_PARTS = [
    "word/document.xml",
    "word/footnotes.xml",
    "word/endnotes.xml",
]
WORD_RE = re.compile(r"[A-Za-zÀ-ÖØ-öø-ÿ0-9][A-Za-zÀ-ÖØ-öø-ÿ0-9'/-]*")
DUTCH_STOPWORDS = {
    "aan",
    "als",
    "bij",
    "dat",
    "de",
    "deze",
    "die",
    "dit",
    "door",
    "een",
    "en",
    "er",
    "het",
    "hun",
    "in",
    "is",
    "met",
    "naar",
    "niet",
    "om",
    "of",
    "op",
    "ook",
    "te",
    "tot",
    "uit",
    "van",
    "voor",
    "waar",
    "wij",
    "wordt",
    "ze",
    "zijn",
    "zo",
}
ABSTRACT_TERMS = [
    "adviesgedreven",
    "ambitie",
    "bedrijfsmodel",
    "continuïteit",
    "data-gedreven",
    "ecosysteem",
    "executiekracht",
    "geïntegreerd",
    "governance",
    "groeicapaciteit",
    "informatievoorsprong",
    "innovatie",
    "klantbinding",
    "klantvraagstuk",
    "markttoegang",
    "onderscheidend",
    "onderscheidend vermogen",
    "oplossinggedreven",
    "platform",
    "platformdenken",
    "platformdiensten",
    "platformmodel",
    "proactievere",
    "processturing",
    "propositie",
    "regie",
    "rendement",
    "schaalbaar",
    "strategisch",
    "structureel",
    "terugkerende",
    "toegevoegde waarde",
    "transformatie",
    "verdienmodel",
]
CONSULTING_TERMS = [
    "aanbodgestuurd",
    "bestuurlijke verantwoordelijkheid",
    "commercieel model",
    "financiële discipline",
    "leidende positie",
    "op het snijvlak",
    "probleemgestuurd",
    "strategische relevantie",
    "structureel onderdeel",
    "vertaalslag",
]

try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except AttributeError:
    pass


@dataclass
class Paragraph:
    id: str
    part: str
    index: int
    text: str


class TextHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        if data.strip():
            self.parts.append(data)

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.lower() in {"p", "br", "div", "li", "tr", "h1", "h2", "h3", "h4"}:
            self.parts.append("\n")

    def text(self) -> str:
        raw = " ".join(self.parts)
        raw = html.unescape(raw)
        raw = re.sub(r"[ \t\r\f\v]+", " ", raw)
        raw = re.sub(r"\n\s+", "\n", raw)
        return raw.strip()


def qn(local_name: str) -> str:
    return f"{{{W_NS}}}{local_name}"


def require_lxml() -> Any:
    if LET is None:
        raise RuntimeError("DOCX replacement requires lxml to preserve Word XML namespaces.")
    return LET


def iter_docx_parts(path: Path) -> Iterable[tuple[str, ET.Element]]:
    with zipfile.ZipFile(path) as zf:
        for part in WORD_PARTS:
            try:
                yield part, ET.fromstring(zf.read(part))
            except KeyError:
                continue


def paragraph_text(paragraph: ET.Element) -> str:
    chunks: list[str] = []
    for node in paragraph.iter():
        if node.tag == qn("t"):
            chunks.append(node.text or "")
        elif node.tag == qn("tab"):
            chunks.append("\t")
        elif node.tag == qn("br"):
            chunks.append("\n")
    return "".join(chunks)


def paragraph_prefix(part: str) -> str:
    if part == "word/document.xml":
        return "body"
    return Path(part).stem


def docx_paragraphs(path: Path) -> list[Paragraph]:
    paragraphs: list[Paragraph] = []
    counters: dict[str, int] = {}
    for part, root in iter_docx_parts(path):
        prefix = paragraph_prefix(part)
        counters.setdefault(prefix, 0)
        for paragraph in root.findall(".//w:p", NS):
            text = paragraph_text(paragraph)
            para_id = f"{prefix}:p{counters[prefix]:06d}"
            counters[prefix] += 1
            paragraphs.append(Paragraph(para_id, part, counters[prefix] - 1, text))
    return paragraphs


def extract_docx(path: Path) -> list[dict[str, Any]]:
    return [
        {"id": p.id, "part": p.part, "index": p.index, "text": p.text}
        for p in docx_paragraphs(path)
        if p.text.strip()
    ]


def extract_pdf(path: Path) -> str:
    try:
        import pypdf  # type: ignore

        reader = pypdf.PdfReader(str(path))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages).strip()
    except ImportError:
        pass

    try:
        import PyPDF2  # type: ignore

        reader = PyPDF2.PdfReader(str(path))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages).strip()
    except ImportError:
        pass

    try:
        import pdfplumber  # type: ignore

        with pdfplumber.open(str(path)) as pdf:
            return "\n\n".join(page.extract_text() or "" for page in pdf.pages).strip()
    except ImportError:
        pass

    try:
        import fitz  # type: ignore

        doc = fitz.open(str(path))
        return "\n\n".join(page.get_text("text") for page in doc).strip()
    except ImportError:
        pass

    raise RuntimeError(
        "PDF extraction requires one installed backend: pypdf, PyPDF2, pdfplumber, or PyMuPDF."
    )


def extract_textlike(path: Path) -> str:
    raw = path.read_text(encoding="utf-8-sig", errors="replace")
    if path.suffix.lower() in {".html", ".htm"}:
        parser = TextHTMLParser()
        parser.feed(raw)
        return parser.text()
    return raw


def extract_any(path: Path) -> list[dict[str, Any]]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        text = extract_pdf(path)
        return [{"id": "pdf:p000000", "part": str(path), "index": 0, "text": text}]
    if suffix in {".txt", ".md", ".markdown", ".html", ".htm"}:
        return [{"id": "text:p000000", "part": str(path), "index": 0, "text": extract_textlike(path)}]
    raise RuntimeError(f"Unsupported input type: {path.suffix}")


def write_extraction(items: list[dict[str, Any]], output_format: str) -> None:
    if output_format == "json":
        json.dump({"paragraphs": items}, sys.stdout, ensure_ascii=False, indent=2)
        sys.stdout.write("\n")
        return
    for item in items:
        text = item["text"].strip()
        if text:
            print(f"[{item['id']}] {text}")


def set_paragraph_text(paragraph: ET.Element, text: str) -> None:
    etree = require_lxml()

    def write_text_nodes(run: ET.Element, value: str) -> None:
        parts = value.split("\n")
        for index, part in enumerate(parts):
            if index:
                etree.SubElement(run, qn("br"))
            t_node = etree.SubElement(run, qn("t"))
            t_node.text = part
            if part.startswith(" ") or part.endswith(" "):
                t_node.set(XML_SPACE, "preserve")

    runs = paragraph.findall("w:r", NS)
    if not runs:
        run = etree.SubElement(paragraph, qn("r"))
        write_text_nodes(run, text)
        return

    first_run = runs[0]
    for run in runs:
        for child in list(run):
            if child.tag in {qn("t"), qn("tab"), qn("br")}:
                run.remove(child)

    write_text_nodes(first_run, text)


def load_replacements(path: Path) -> dict[str, str]:
    data = json.loads(path.read_text(encoding="utf-8"))
    rows = data.get("paragraphs")
    if not isinstance(rows, list):
        raise RuntimeError("Replacement JSON must contain a 'paragraphs' array.")

    replacements: dict[str, str] = {}
    for row in rows:
        if not isinstance(row, dict) or "id" not in row or "text" not in row:
            raise RuntimeError("Each replacement row must contain 'id' and 'text'.")
        replacements[str(row["id"])] = str(row["text"])
    return replacements


def xml_standalone(raw: bytes) -> bool | None:
    prefix = raw[:200].lower()
    if b"standalone" not in prefix:
        return None
    if b"standalone='yes'" in prefix or b'standalone="yes"' in prefix:
        return True
    if b"standalone='no'" in prefix or b'standalone="no"' in prefix:
        return False
    return None


def parse_word_xml(raw: bytes) -> Any:
    etree = require_lxml()
    parser = etree.XMLParser(remove_blank_text=False, resolve_entities=False, huge_tree=True)
    return etree.fromstring(raw, parser)


def serialize_word_xml(root: Any, original_raw: bytes) -> bytes:
    etree = require_lxml()
    kwargs: dict[str, Any] = {
        "encoding": "UTF-8",
        "xml_declaration": True,
        "pretty_print": False,
    }
    standalone = xml_standalone(original_raw)
    if standalone is not None:
        kwargs["standalone"] = standalone
    return etree.tostring(root, **kwargs)


def replace_docx(source: Path, replacements_path: Path, output: Path) -> int:
    replacements = load_replacements(replacements_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    modified_parts: dict[str, bytes] = {}
    with zipfile.ZipFile(source) as zf:
        replaced = 0
        applied_ids: set[str] = set()
        for part in WORD_PARTS:
            try:
                original_raw = zf.read(part)
            except KeyError:
                continue
            doc_root = parse_word_xml(original_raw)
            prefix = paragraph_prefix(part)
            part_replaced = 0
            for index, paragraph in enumerate(doc_root.findall(".//w:p", NS)):
                para_id = f"{prefix}:p{index:06d}"
                if para_id in replacements:
                    set_paragraph_text(paragraph, replacements[para_id])
                    replaced += 1
                    part_replaced += 1
                    applied_ids.add(para_id)
            if part_replaced:
                modified_parts[part] = serialize_word_xml(doc_root, original_raw)

        missing_ids = sorted(set(replacements) - applied_ids)
        if missing_ids:
            preview = ", ".join(missing_ids[:10])
            suffix = "..." if len(missing_ids) > 10 else ""
            raise RuntimeError(f"Replacement IDs not found in source DOCX: {preview}{suffix}")

        tmp_fd, tmp_name = tempfile.mkstemp(
            prefix=f"{output.stem}.", suffix=".tmp.docx", dir=str(output.parent)
        )
        os.close(tmp_fd)
        tmp_path = Path(tmp_name)
        try:
            with zipfile.ZipFile(tmp_path, "w") as out_zf:
                for info in zf.infolist():
                    data = modified_parts.get(info.filename)
                    if data is None:
                        data = zf.read(info.filename)
                    out_zf.writestr(info, data)
            validate_docx_package(tmp_path, source=source, fail_on_text=False)
            os.replace(tmp_path, output)
        finally:
            if tmp_path.exists():
                tmp_path.unlink()
    return replaced


def docx_entry_names(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as zf:
        return zf.namelist()


def validate_docx_package(path: Path, source: Path | None = None, fail_on_text: bool = True) -> tuple[list[str], list[str]]:
    failures: list[str] = []
    warnings: list[str] = []
    required = {"[Content_Types].xml", "_rels/.rels", "word/document.xml", "word/_rels/document.xml.rels"}

    try:
        with zipfile.ZipFile(path) as zf:
            bad = zf.testzip()
            if bad:
                failures.append(f"ZIP CRC failure at {bad}")
            names = zf.namelist()
            missing = sorted(required - set(names))
            if missing:
                failures.append(f"Missing required DOCX parts: {', '.join(missing)}")
            for part in WORD_PARTS:
                if part in names:
                    raw = zf.read(part)
                    try:
                        parse_word_xml(raw)
                    except Exception as exc:  # noqa: BLE001
                        failures.append(f"{part} does not parse as XML: {exc}")
            document_xml = zf.read("word/document.xml")
            if b"<ns0:document" in document_xml[:500] or b"ns0:" in document_xml[:500]:
                failures.append("word/document.xml uses generic ns0 prefixes; Word may treat this as repairable content.")
            if b"<w:document" not in document_xml[:500]:
                warnings.append("word/document.xml does not start with a w:document prefix.")
    except zipfile.BadZipFile as exc:
        failures.append(f"Not a valid ZIP/DOCX package: {exc}")

    if source is not None:
        try:
            source_names = docx_entry_names(source)
            target_names = docx_entry_names(path)
            if set(source_names) != set(target_names):
                failures.append("DOCX package entries differ from source.")
            if source_names != target_names:
                warnings.append("DOCX package entry order differs from source.")
        except Exception as exc:  # noqa: BLE001
            failures.append(f"Could not compare package entries with source: {exc}")

    try:
        import docx  # type: ignore

        docx.Document(str(path))
    except ImportError:
        warnings.append("python-docx is not installed; skipped python-docx open check.")
    except Exception as exc:  # noqa: BLE001
        failures.append(f"python-docx could not open document: {exc}")

    if fail_on_text:
        try:
            items = extract_any(path)
            findings = audit_text("\n\n".join(item["text"] for item in items))
            if findings:
                failures.extend(f"text audit: {finding}" for finding in findings)
        except Exception as exc:  # noqa: BLE001
            failures.append(f"Could not run text audit: {exc}")

    return failures, warnings


def audit_text(text: str) -> list[str]:
    checks: list[tuple[str, str]] = [
        ("em dash", "—"),
        ("en dash", "–"),
        ("curly double quote", r"[“”]"),
        ("curly single quote", r"[‘’]"),
        ("emoji/symbol bullet", r"[\U0001F300-\U0001FAFF]"),
        ("bold inline header", r"\*\*[^*\n]{1,60}:\*\*"),
        ("chatbot residue", r"\b(natuurlijk|zeker|ik hoop dat dit helpt|laat het me weten|laten we|hier is een overzicht)\b"),
        ("inflated Dutch significance", r"\b(cruciaal|essentieel|pivotaal|sleutelrol|onderstreept|toonaangevend|baanbrekend|naadloos|robuust|dynamisch)\b"),
        ("vague authority", r"\b(experts stellen|onderzoek toont aan|diverse bronnen|marktpartijen zien|men verwacht)\b"),
        ("English AI vocabulary", r"\b(delven|delve|showcase|landscape|tapestry|testament|pivotal|crucial|unlocking potential)\b"),
        ("fake contrast", r"\b(niet alleen|meer dan alleen|not only|not just)\b"),
    ]
    findings: list[str] = []
    for label, pattern in checks:
        if label in {"em dash", "en dash"}:
            count = text.count(pattern)
        else:
            count = len(re.findall(pattern, text, flags=re.IGNORECASE))
        if count:
            findings.append(f"{label}: {count}")
    return findings


def words_in(text: str) -> list[str]:
    return [match.group(0).lower() for match in WORD_RE.finditer(text)]


def content_words_in(text: str) -> list[str]:
    return [
        word
        for word in words_in(text)
        if len(word) > 3 and word not in DUTCH_STOPWORDS and not word.isdigit()
    ]


def split_sentences(text: str) -> list[str]:
    sentences: list[str] = []
    for line in re.split(r"\n+", text):
        cleaned = line.strip().strip("•- \t")
        if not cleaned:
            continue
        parts = re.split(r"(?<=[.!?])\s+(?=[A-ZÀ-Ö0-9\"'])", cleaned)
        sentences.extend(part.strip() for part in parts if part.strip())
    return sentences


def series_stats(values: list[int]) -> dict[str, Any]:
    if not values:
        return {"count": 0, "average": 0, "stdev": 0, "coefficient_of_variation": 0, "min": 0, "max": 0}
    avg = mean(values)
    deviation = pstdev(values) if len(values) > 1 else 0
    return {
        "count": len(values),
        "average": round(avg, 2),
        "stdev": round(deviation, 2),
        "coefficient_of_variation": round(deviation / avg, 2) if avg else 0,
        "min": min(values),
        "max": max(values),
    }


def count_terms(text: str, terms: list[str]) -> dict[str, int]:
    lower = text.lower()
    counts: dict[str, int] = {}
    for term in terms:
        pattern = r"(?<![\wÀ-ÖØ-öø-ÿ])" + re.escape(term.lower()) + r"(?![\wÀ-ÖØ-öø-ÿ])"
        count = len(re.findall(pattern, lower, flags=re.IGNORECASE))
        if count:
            counts[term] = count
    return counts


def top_repeated_words(text: str, limit: int = 12) -> list[dict[str, Any]]:
    counts = Counter(content_words_in(text))
    return [
        {"word": word, "count": count}
        for word, count in counts.most_common(limit)
        if count > 1
    ]


def top_repeated_ngrams(text: str, n: int, limit: int = 10) -> list[dict[str, Any]]:
    words = words_in(text)
    counts: Counter[str] = Counter()
    abstract_words = {term for term in ABSTRACT_TERMS if " " not in term}
    for index in range(0, max(0, len(words) - n + 1)):
        chunk = words[index : index + n]
        if all(word in DUTCH_STOPWORDS for word in chunk):
            continue
        if any(len(word) <= 1 for word in chunk):
            continue
        if chunk[0] in DUTCH_STOPWORDS or chunk[-1] in DUTCH_STOPWORDS:
            allowed_article_anchor = (
                n == 2
                and chunk[0] in {"de", "het", "een"}
                and chunk[1] in abstract_words
            )
            if not allowed_article_anchor:
                continue
        counts[" ".join(chunk)] += 1
    return [
        {"phrase": phrase, "count": count}
        for phrase, count in counts.most_common(limit)
        if count > 1
    ]


def concrete_anchor_counts(text: str) -> dict[str, int]:
    numbers = len(re.findall(r"\b\d+(?:[.,]\d+)?\b", text))
    currency = len(re.findall(r"€\s?\d+(?:[.,]\d+)?", text))
    percentages = len(re.findall(r"\b\d+(?:[.,]\d+)?\s?%", text))
    entities = len(
        re.findall(
            r"\b[A-ZÀ-Ý][A-Za-zÀ-ÖØ-öø-ÿ]*(?:[ /&.-]+[A-ZÀ-Ý][A-Za-zÀ-ÖØ-öø-ÿ]*)*\b",
            text,
        )
    )
    return {
        "numbers": numbers,
        "currency": currency,
        "percentages": percentages,
        "capitalized_entities": entities,
        "total": numbers + currency + percentages + entities,
    }


def section_structure(paragraphs: list[str]) -> dict[str, Any]:
    headings: list[str] = []
    bullet_lines = 0
    for paragraph in paragraphs:
        lines = [line.strip() for line in paragraph.splitlines() if line.strip()]
        if not lines:
            continue
        bullet_lines += sum(1 for line in lines if line.startswith(("•", "- ")))
        first = lines[0]
        if len(lines) > 1 and len(first) <= 90 and not first.endswith((".", ":", ";", ",")):
            headings.append(first)
    heading_ratio = len(headings) / len(paragraphs) if paragraphs else 0
    return {
        "heading_body_sections": len(headings),
        "heading_body_ratio": round(heading_ratio, 2),
        "bullet_lines": bullet_lines,
        "headings": headings,
    }


def score_profile(
    word_count: int,
    abstract_hit_count: int,
    consulting_hit_count: int,
    repeated_bigram_count: int,
    paragraph_stats: dict[str, Any],
    sentence_stats: dict[str, Any],
    structure: dict[str, Any],
    concrete_total: int,
) -> dict[str, Any]:
    if word_count <= 0:
        return {"score": 0, "label": "empty", "signals": []}

    abstract_density = abstract_hit_count / word_count * 100
    concrete_density = concrete_total / word_count * 100
    score = 0
    signals: list[str] = []

    if abstract_density >= 7:
        score += 30
        signals.append("high abstract Dutch business vocabulary density")
    elif abstract_density >= 4:
        score += 18
        signals.append("moderate abstract Dutch business vocabulary density")

    if consulting_hit_count >= 6:
        score += 18
        signals.append("many consulting-style stock phrases")
    elif consulting_hit_count >= 3:
        score += 10
        signals.append("some consulting-style stock phrases")

    if repeated_bigram_count >= 3:
        score += 18
        signals.append("many repeated two-word phrases")
    elif repeated_bigram_count >= 2:
        score += 10
        signals.append("some repeated two-word phrases")

    if structure["heading_body_ratio"] >= 0.65 and structure["heading_body_sections"] >= 6:
        score += 18
        signals.append("highly regular heading-plus-paragraph section skeleton")

    if paragraph_stats["count"] >= 8 and paragraph_stats["coefficient_of_variation"] < 0.5:
        score += 8
        signals.append("paragraph lengths are unusually even")
    if sentence_stats["count"] >= 12 and sentence_stats["coefficient_of_variation"] < 0.45:
        score += 6
        signals.append("sentence lengths are unusually even")

    if concrete_density >= 10:
        score -= 8
    elif concrete_density < 4 and word_count >= 250:
        score += 8
        signals.append("few concrete anchors compared with document length")

    score = max(0, min(100, score))
    if score >= 65:
        label = "high AI-like structural risk"
    elif score >= 35:
        label = "moderate AI-like structural risk"
    else:
        label = "low AI-like structural risk"
    return {"score": score, "label": label, "signals": signals}


def humanity_profile(items: list[dict[str, Any]]) -> dict[str, Any]:
    paragraphs = [str(item["text"]).strip() for item in items if str(item["text"]).strip()]
    text = "\n\n".join(paragraphs)
    words = words_in(text)
    sentences = split_sentences(text)
    abstract_counts = count_terms(text, ABSTRACT_TERMS)
    consulting_counts = count_terms(text, CONSULTING_TERMS)
    repeated_bigrams = top_repeated_ngrams(text, 2)
    repeated_trigrams = top_repeated_ngrams(text, 3)
    significant_repeated_bigrams = sum(1 for row in repeated_bigrams if row["count"] >= 3)
    concrete = concrete_anchor_counts(text)
    paragraph_stats = series_stats([len(words_in(paragraph)) for paragraph in paragraphs])
    sentence_stats = series_stats([len(words_in(sentence)) for sentence in sentences])
    structure = section_structure(paragraphs)
    abstract_hit_count = sum(abstract_counts.values())
    consulting_hit_count = sum(consulting_counts.values())
    score = score_profile(
        word_count=len(words),
        abstract_hit_count=abstract_hit_count,
        consulting_hit_count=consulting_hit_count,
        repeated_bigram_count=significant_repeated_bigrams,
        paragraph_stats=paragraph_stats,
        sentence_stats=sentence_stats,
        structure=structure,
        concrete_total=concrete["total"],
    )
    return {
        "counts": {
            "paragraphs": len(paragraphs),
            "sentences": len(sentences),
            "words": len(words),
        },
        "paragraph_words": paragraph_stats,
        "sentence_words": sentence_stats,
        "abstract_terms": {
            "total": abstract_hit_count,
            "density_per_100_words": round(abstract_hit_count / len(words) * 100, 2) if words else 0,
            "terms": abstract_counts,
        },
        "consulting_terms": {
            "total": consulting_hit_count,
            "density_per_100_words": round(consulting_hit_count / len(words) * 100, 2) if words else 0,
            "terms": consulting_counts,
        },
        "concrete_anchors": {
            **concrete,
            "density_per_100_words": round(concrete["total"] / len(words) * 100, 2) if words else 0,
        },
        "repetition": {
            "words": top_repeated_words(text),
            "bigrams": repeated_bigrams,
            "trigrams": repeated_trigrams,
            "significant_bigram_count": significant_repeated_bigrams,
        },
        "structure": structure,
        "obvious_ai_tells": audit_text(text),
        "risk": score,
    }


def print_humanity_profile(profile: dict[str, Any]) -> None:
    counts = profile["counts"]
    print(f"Risk: {profile['risk']['score']} ({profile['risk']['label']})")
    print(
        "Counts: "
        f"{counts['paragraphs']} paragraphs, {counts['sentences']} sentences, {counts['words']} words"
    )
    print(f"Paragraph words: {profile['paragraph_words']}")
    print(f"Sentence words: {profile['sentence_words']}")
    print(
        "Abstract terms: "
        f"{profile['abstract_terms']['total']} "
        f"({profile['abstract_terms']['density_per_100_words']} per 100 words)"
    )
    print(
        "Consulting terms: "
        f"{profile['consulting_terms']['total']} "
        f"({profile['consulting_terms']['density_per_100_words']} per 100 words)"
    )
    print(
        "Concrete anchors: "
        f"{profile['concrete_anchors']['total']} "
        f"({profile['concrete_anchors']['density_per_100_words']} per 100 words)"
    )
    print(
        "Structure: "
        f"{profile['structure']['heading_body_sections']} heading-body sections, "
        f"ratio {profile['structure']['heading_body_ratio']}, "
        f"{profile['structure']['bullet_lines']} bullet lines"
    )
    if profile["risk"]["signals"]:
        print("Signals:")
        for signal in profile["risk"]["signals"]:
            print(f"- {signal}")
    if profile["repetition"]["words"]:
        print("Repeated words:")
        for row in profile["repetition"]["words"][:8]:
            print(f"- {row['word']}: {row['count']}")
    if profile["repetition"]["bigrams"]:
        print(f"Repeated phrases (3+ count: {profile['repetition']['significant_bigram_count']}):")
        for row in profile["repetition"]["bigrams"][:8]:
            print(f"- {row['phrase']}: {row['count']}")
    if profile["obvious_ai_tells"]:
        print("Obvious tells:")
        for finding in profile["obvious_ai_tells"]:
            print(f"- {finding}")


def read_rewrite_text(path_arg: str) -> str:
    if path_arg == "-":
        raw = sys.stdin.buffer.read()
        try:
            return raw.decode("utf-8-sig")
        except UnicodeDecodeError:
            return raw.decode(sys.stdin.encoding or "utf-8", errors="replace")
    return Path(path_arg).read_text(encoding="utf-8-sig")


def add_docx_block(doc: Any, block: str, block_index: int) -> None:
    lines = [line.rstrip() for line in block.splitlines() if line.strip()]
    if not lines:
        return

    first = lines[0].strip()
    if first.startswith("# "):
        doc.add_heading(first[2:].strip(), level=1)
        return
    if first.startswith("## "):
        doc.add_heading(first[3:].strip(), level=2)
        return
    if first.startswith("### "):
        doc.add_heading(first[4:].strip(), level=3)
        return

    if all(line.strip().startswith(("- ", "• ")) for line in lines):
        for line in lines:
            text = line.strip()[2:].strip()
            if text:
                doc.add_paragraph(text, style="List Bullet")
        return

    if block_index == 0 and len(lines) == 1 and len(first) <= 100 and not first.endswith("."):
        doc.add_paragraph(first, style="Title")
        return

    paragraph = doc.add_paragraph()
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        paragraph.add_run(line.strip())


def write_rebuilt_docx(source: Path | None, text: str, output: Path) -> None:
    try:
        import docx  # type: ignore
        from docx.shared import Cm, Pt  # type: ignore
    except ImportError as exc:  # pragma: no cover - depends on local environment.
        raise RuntimeError("write-docx-from-text requires python-docx.") from exc

    if source is not None and source.suffix.lower() != ".docx":
        raise RuntimeError("Source must be a DOCX file.")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(20)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(12)

    blocks = [block.strip() for block in re.split(r"\n\s*\n", text.strip()) if block.strip()]
    if not blocks:
        raise RuntimeError("No text blocks found for DOCX output.")

    for index, block in enumerate(blocks):
        add_docx_block(doc, block, index)

    if source is not None:
        doc.core_properties.comments = f"Rebuilt from {source.name} by dutch-text-humanizer."
    else:
        doc.core_properties.comments = "Created by dutch-text-humanizer."
    doc.save(str(output))
    failures, _warnings = validate_docx_package(output, source=None, fail_on_text=False)
    if failures:
        raise RuntimeError("DOCX package validation failed after rebuild: " + "; ".join(failures))


def supported_corpus_files(path: Path) -> list[Path]:
    supported = {".docx", ".pdf", ".txt", ".md", ".markdown", ".html", ".htm"}
    if path.is_file():
        if path.suffix.lower() not in supported:
            raise RuntimeError(f"Unsupported corpus file type: {path.suffix}")
        return [path]
    if not path.is_dir():
        raise RuntimeError(f"Corpus path not found: {path}")
    files = [
        item
        for item in sorted(path.rglob("*"))
        if item.is_file()
        and item.suffix.lower() in supported
        and not item.name.startswith("~$")
    ]
    if not files:
        raise RuntimeError(f"No supported corpus files found in {path}")
    return files


def profile_corpus(path: Path) -> dict[str, Any]:
    files = supported_corpus_files(path)
    aggregate_items: list[dict[str, Any]] = []
    file_profiles: list[dict[str, Any]] = []
    for file_path in files:
        items = extract_any(file_path)
        profile = humanity_profile(items)
        file_profiles.append(
            {
                "path": str(file_path),
                "profile": profile,
            }
        )
        for item in items:
            aggregate_items.append(
                {
                    "id": f"{file_path.name}:{item['id']}",
                    "part": item.get("part", str(file_path)),
                    "index": item.get("index", 0),
                    "text": item["text"],
                }
            )
    return {
        "path": str(path),
        "file_count": len(files),
        "aggregate": humanity_profile(aggregate_items),
        "files": file_profiles,
    }


def profile_metric_summary(profile: dict[str, Any]) -> dict[str, float]:
    return {
        "risk_score": float(profile["risk"]["score"]),
        "abstract_density": float(profile["abstract_terms"]["density_per_100_words"]),
        "consulting_density": float(profile["consulting_terms"]["density_per_100_words"]),
        "concrete_density": float(profile["concrete_anchors"]["density_per_100_words"]),
        "paragraph_cv": float(profile["paragraph_words"]["coefficient_of_variation"]),
        "sentence_cv": float(profile["sentence_words"]["coefficient_of_variation"]),
        "heading_body_ratio": float(profile["structure"]["heading_body_ratio"]),
        "significant_repeated_bigrams": float(profile["repetition"]["significant_bigram_count"]),
    }


def style_comparison(corpus_path: Path, target_path: Path) -> dict[str, Any]:
    corpus = profile_corpus(corpus_path)
    target = humanity_profile(extract_any(target_path))
    corpus_metrics = profile_metric_summary(corpus["aggregate"])
    target_metrics = profile_metric_summary(target)
    deltas = {
        key: round(target_metrics[key] - corpus_metrics[key], 2)
        for key in corpus_metrics
    }
    penalties = [
        min(1.0, max(0.0, deltas["abstract_density"]) / 3),
        min(1.0, max(0.0, deltas["consulting_density"]) / 1),
        min(1.0, max(0.0, -deltas["concrete_density"]) / 8),
        min(1.0, max(0.0, -deltas["paragraph_cv"]) / 2),
        min(1.0, max(0.0, -deltas["sentence_cv"]) / 1),
        min(1.0, max(0.0, deltas["heading_body_ratio"]) / 0.6),
        min(1.0, max(0.0, deltas["significant_repeated_bigrams"]) / 5),
    ]
    distance = round(sum(penalties) / len(penalties) * 100, 1)
    match_score = round(max(0.0, 100 - distance), 1)
    if match_score >= 75:
        label = "strong natural-style match"
    elif match_score >= 50:
        label = "partial natural-style match"
    else:
        label = "weak natural-style match"
    return {
        "corpus": corpus,
        "target": {
            "path": str(target_path),
            "profile": target,
        },
        "metrics": {
            "corpus": corpus_metrics,
            "target": target_metrics,
            "delta": deltas,
        },
        "style_match": {
            "score": match_score,
            "distance": distance,
            "label": label,
        },
    }


def print_corpus_profile(payload: dict[str, Any]) -> None:
    print(f"Corpus: {payload['path']}")
    print(f"Files: {payload['file_count']}")
    print("\nAggregate:")
    print_humanity_profile(payload["aggregate"])
    print("\nFiles:")
    for row in payload["files"]:
        profile = row["profile"]
        print(
            f"- {Path(row['path']).name}: "
            f"{profile['counts']['words']} words, "
            f"risk {profile['risk']['score']}, "
            f"abstract {profile['abstract_terms']['density_per_100_words']} per 100, "
            f"paragraph CV {profile['paragraph_words']['coefficient_of_variation']}"
        )


def print_style_comparison(payload: dict[str, Any]) -> None:
    match = payload["style_match"]
    print(f"Style match: {match['score']} ({match['label']})")
    print(f"Distance: {match['distance']}")
    print("\nMetric deltas (target - corpus):")
    for key, value in payload["metrics"]["delta"].items():
        print(f"- {key}: {value:+}")
    print("\nTarget profile:")
    print_humanity_profile(payload["target"]["profile"])


def command_extract(args: argparse.Namespace) -> None:
    items = extract_any(Path(args.input))
    write_extraction(items, args.format)


def command_map_docx(args: argparse.Namespace) -> None:
    path = Path(args.input)
    items = [
        {"id": p.id, "part": p.part, "index": p.index, "text": p.text}
        for p in docx_paragraphs(path)
    ]
    payload = {"paragraphs": items}
    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {len(items)} paragraphs to {out}")


def command_replace_docx(args: argparse.Namespace) -> None:
    count = replace_docx(Path(args.source), Path(args.replacements), Path(args.output))
    print(f"Replaced {count} paragraphs in {args.output}")


def command_audit(args: argparse.Namespace) -> None:
    items = extract_any(Path(args.input))
    text = "\n\n".join(item["text"] for item in items)
    findings = audit_text(text)
    if findings:
        print("Findings:")
        for finding in findings:
            print(f"- {finding}")
        sys.exit(1)
    print("No configured AI-writing tells found.")


def command_audit_docx(args: argparse.Namespace) -> None:
    failures, warnings = validate_docx_package(
        Path(args.input), source=Path(args.source) if args.source else None
    )
    if warnings:
        print("Warnings:")
        for warning in warnings:
            print(f"- {warning}")
    if failures:
        print("Failures:")
        for failure in failures:
            print(f"- {failure}")
        sys.exit(1)
    print("DOCX package audit passed.")


def command_profile_humanity(args: argparse.Namespace) -> None:
    profile = humanity_profile(extract_any(Path(args.input)))
    if args.format == "json":
        json.dump(profile, sys.stdout, ensure_ascii=False, indent=2)
        sys.stdout.write("\n")
        return
    print_humanity_profile(profile)


def command_compare_humanity(args: argparse.Namespace) -> None:
    before = humanity_profile(extract_any(Path(args.before)))
    after = humanity_profile(extract_any(Path(args.after)))
    payload = {
        "before": before,
        "after": after,
        "delta": {
            "risk_score": after["risk"]["score"] - before["risk"]["score"],
            "abstract_terms": after["abstract_terms"]["total"] - before["abstract_terms"]["total"],
            "consulting_terms": after["consulting_terms"]["total"] - before["consulting_terms"]["total"],
            "significant_repeated_bigrams": (
                after["repetition"]["significant_bigram_count"]
                - before["repetition"]["significant_bigram_count"]
            ),
            "heading_body_sections": (
                after["structure"]["heading_body_sections"] - before["structure"]["heading_body_sections"]
            ),
            "words": after["counts"]["words"] - before["counts"]["words"],
        },
    }
    if args.format == "json":
        json.dump(payload, sys.stdout, ensure_ascii=False, indent=2)
        sys.stdout.write("\n")
        return

    print("Before:")
    print_humanity_profile(before)
    print("\nAfter:")
    print_humanity_profile(after)
    print("\nDelta:")
    for key, value in payload["delta"].items():
        print(f"- {key}: {value:+}")


def command_profile_corpus(args: argparse.Namespace) -> None:
    payload = profile_corpus(Path(args.corpus))
    if args.format == "json":
        json.dump(payload, sys.stdout, ensure_ascii=False, indent=2)
        sys.stdout.write("\n")
        return
    print_corpus_profile(payload)


def command_compare_style(args: argparse.Namespace) -> None:
    payload = style_comparison(Path(args.corpus), Path(args.target))
    if args.format == "json":
        json.dump(payload, sys.stdout, ensure_ascii=False, indent=2)
        sys.stdout.write("\n")
        return
    print_style_comparison(payload)


def command_write_docx_from_text(args: argparse.Namespace) -> None:
    text = read_rewrite_text(args.text)
    write_rebuilt_docx(Path(args.source), text, Path(args.output))
    print(f"Wrote rebuilt DOCX to {args.output}")


def command_create_docx_from_text(args: argparse.Namespace) -> None:
    text = read_rewrite_text(args.text)
    write_rebuilt_docx(None, text, Path(args.output))
    print(f"Wrote DOCX to {args.output}")


def command_diff_docx(args: argparse.Namespace) -> None:
    before = {item["id"]: item for item in extract_docx(Path(args.source))}
    after = {item["id"]: item for item in extract_docx(Path(args.target))}
    all_ids = sorted(set(before) | set(after))
    changes: list[dict[str, Any]] = []
    for para_id in all_ids:
        old = before.get(para_id, {}).get("text")
        new = after.get(para_id, {}).get("text")
        if old != new:
            changes.append({"id": para_id, "before": old, "after": new})

    if args.format == "json":
        payload = {"changed_count": len(changes), "changes": changes}
        if args.out:
            Path(args.out).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        else:
            json.dump(payload, sys.stdout, ensure_ascii=False, indent=2)
            sys.stdout.write("\n")
        return

    print(f"Changed paragraphs: {len(changes)}")
    for change in changes:
        print(f"[{change['id']}]")
        print(f"- {change['before']}")
        print(f"+ {change['after']}")


def command_visualize_docx(args: argparse.Namespace) -> None:
    input_path = Path(args.input)
    out_dir = Path(args.output_dir)
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    winword = shutil.which("WINWORD")

    try:
        import pdf2image  # noqa: F401

        pdf2image_available = True
    except ImportError:
        pdf2image_available = False

    report = {
        "input": str(input_path),
        "output_dir": str(out_dir),
        "soffice": soffice,
        "winword": winword,
        "pdf2image": pdf2image_available,
        "status": "unavailable",
        "message": "",
    }

    if not soffice:
        report["message"] = "No LibreOffice/soffice executable found; DOCX visual rendering is unavailable."
    elif not pdf2image_available:
        report["message"] = "pdf2image is not installed; PDF-to-PNG rasterization is unavailable."
    else:
        out_dir.mkdir(parents=True, exist_ok=True)
        pdf_dir = out_dir / "pdf"
        pdf_dir.mkdir(parents=True, exist_ok=True)
        cmd = [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_dir),
            str(input_path),
        ]
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=args.timeout)
        if proc.returncode != 0:
            report["message"] = f"LibreOffice conversion failed: {proc.stderr.strip() or proc.stdout.strip()}"
        else:
            from pdf2image import convert_from_path  # type: ignore

            pdf_candidates = sorted(pdf_dir.glob("*.pdf"))
            if not pdf_candidates:
                report["message"] = "LibreOffice completed but did not produce a PDF."
            else:
                pages = convert_from_path(str(pdf_candidates[0]), dpi=args.dpi)
                png_dir = out_dir / "png"
                png_dir.mkdir(parents=True, exist_ok=True)
                for index, page in enumerate(pages, start=1):
                    page.save(png_dir / f"page-{index:03d}.png")
                report["status"] = "completed"
                report["message"] = f"Rendered {len(pages)} pages."
                report["png_dir"] = str(png_dir)

    if args.format == "json":
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        print(f"Status: {report['status']}")
        print(report["message"])
        if report.get("png_dir"):
            print(f"PNG output: {report['png_dir']}")
    if report["status"] != "completed":
        sys.exit(1)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    subparsers = parser.add_subparsers(dest="command", required=True)

    extract_parser = subparsers.add_parser("extract", help="Extract text from DOCX, PDF, text, Markdown, or HTML.")
    extract_parser.add_argument("input")
    extract_parser.add_argument("--format", choices=["text", "json"], default="text")
    extract_parser.set_defaults(func=command_extract)

    map_parser = subparsers.add_parser("map-docx", help="Write DOCX paragraph IDs and text to JSON.")
    map_parser.add_argument("input")
    map_parser.add_argument("--out", required=True)
    map_parser.set_defaults(func=command_map_docx)

    replace_parser = subparsers.add_parser("replace-docx", help="Clone a DOCX and replace paragraphs by ID.")
    replace_parser.add_argument("source")
    replace_parser.add_argument("replacements")
    replace_parser.add_argument("output")
    replace_parser.set_defaults(func=command_replace_docx)

    audit_parser = subparsers.add_parser("audit", help="Scan extracted text for configured AI-writing tells.")
    audit_parser.add_argument("input")
    audit_parser.set_defaults(func=command_audit)

    audit_docx_parser = subparsers.add_parser(
        "audit-docx", help="Validate DOCX package health and run the text audit."
    )
    audit_docx_parser.add_argument("input")
    audit_docx_parser.add_argument("--source", help="Optional source DOCX for package entry comparison.")
    audit_docx_parser.set_defaults(func=command_audit_docx)

    profile_parser = subparsers.add_parser(
        "profile-humanity", help="Profile document-level AI-like structure and Dutch business prose signals."
    )
    profile_parser.add_argument("input")
    profile_parser.add_argument("--format", choices=["text", "json"], default="text")
    profile_parser.set_defaults(func=command_profile_humanity)

    compare_parser = subparsers.add_parser(
        "compare-humanity", help="Compare structural humanization signals between two documents."
    )
    compare_parser.add_argument("before")
    compare_parser.add_argument("after")
    compare_parser.add_argument("--format", choices=["text", "json"], default="text")
    compare_parser.set_defaults(func=command_compare_humanity)

    corpus_parser = subparsers.add_parser(
        "profile-corpus", help="Profile a folder or file corpus for natural-style calibration."
    )
    corpus_parser.add_argument("corpus")
    corpus_parser.add_argument("--format", choices=["text", "json"], default="text")
    corpus_parser.set_defaults(func=command_profile_corpus)

    style_parser = subparsers.add_parser(
        "compare-style", help="Compare a document against a natural-style corpus."
    )
    style_parser.add_argument("corpus")
    style_parser.add_argument("target")
    style_parser.add_argument("--format", choices=["text", "json"], default="text")
    style_parser.set_defaults(func=command_compare_style)

    write_parser = subparsers.add_parser(
        "write-docx-from-text", help="Build a clean DOCX from Markdown-ish rewritten text."
    )
    write_parser.add_argument("source", help="Source DOCX used for provenance and validation context.")
    write_parser.add_argument("text", help="Text file to write, or '-' to read from stdin.")
    write_parser.add_argument("output")
    write_parser.set_defaults(func=command_write_docx_from_text)

    create_parser = subparsers.add_parser(
        "create-docx-from-text", help="Build a clean DOCX from Markdown-ish rewritten text without a source DOCX."
    )
    create_parser.add_argument("text", help="Text file to write, or '-' to read from stdin.")
    create_parser.add_argument("output")
    create_parser.set_defaults(func=command_create_docx_from_text)

    diff_docx_parser = subparsers.add_parser("diff-docx", help="Compare non-empty DOCX paragraphs by ID.")
    diff_docx_parser.add_argument("source")
    diff_docx_parser.add_argument("target")
    diff_docx_parser.add_argument("--format", choices=["text", "json"], default="text")
    diff_docx_parser.add_argument("--out", help="Write JSON diff to this path.")
    diff_docx_parser.set_defaults(func=command_diff_docx)

    visualize_parser = subparsers.add_parser(
        "visualize-docx", help="Render DOCX to page PNGs when local render backends are available."
    )
    visualize_parser.add_argument("input")
    visualize_parser.add_argument("--output-dir", required=True)
    visualize_parser.add_argument("--dpi", type=int, default=150)
    visualize_parser.add_argument("--timeout", type=int, default=120)
    visualize_parser.add_argument("--format", choices=["text", "json"], default="text")
    visualize_parser.set_defaults(func=command_visualize_docx)

    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    try:
        args.func(args)
        return 0
    except Exception as exc:
        print(f"error: {exc}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
